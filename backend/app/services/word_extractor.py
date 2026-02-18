"""Word document (.docx/.doc) text extraction service.

DOCX (Office Open XML): python-docx — paragraphs, tables, headers/footers, embedded images.
DOC (legacy binary): antiword CLI (UTF-8 output, 30s timeout).

Embedded image OCR via OCRService (best-effort, DOCX only — antiword can't extract images).
"""

from __future__ import annotations

import asyncio
import logging
from pathlib import Path

from pydantic import BaseModel

logger = logging.getLogger(__name__)

_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".tif"}

_MIME_MAP = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".bmp": "image/bmp",
    ".gif": "image/gif",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}


class WordExtractionResult(BaseModel):
    """Word document text extraction result."""

    text: str
    page_count: int
    metadata: dict


class WordExtractor:
    """Word document (.docx/.doc) text extraction service."""

    async def extract(self, file_path: str | Path) -> WordExtractionResult:
        """Extract text from a DOCX or DOC file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {path}")

        suffix = path.suffix.lower()
        metadata: dict = {"format": suffix.lstrip(".")}

        if suffix == ".docx":
            text = self._extract_docx(path)
            images = self._extract_docx_images(path)
        elif suffix == ".doc":
            text = await self._extract_doc(path)
            images = []  # antiword can't extract images
        else:
            raise ValueError(f"Unsupported format: {suffix}")

        # OCR embedded images (best-effort)
        ocr_text, ocr_meta = await self._ocr_images(images)
        metadata.update(ocr_meta)

        # Combine body text + OCR text
        parts = [t for t in [text.strip(), ocr_text] if t]
        combined = "\n\n".join(parts)

        if not combined:
            raise ValueError("Word document contains no extractable text")

        return WordExtractionResult(
            text=combined,
            page_count=1,
            metadata=metadata,
        )

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX via python-docx (paragraphs + tables + headers/footers)."""
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []

        # Headers/footers (first section only)
        for section in doc.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
            break  # first section only

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append("\n".join(rows_text))

        # Footers (first section only)
        for section in doc.sections:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
            break

        return "\n".join(parts)

    def _extract_docx_images(self, path: Path) -> list[tuple[str, bytes]]:
        """Extract embedded images from DOCX ZIP structure."""
        import zipfile

        images: list[tuple[str, bytes]] = []
        try:
            with zipfile.ZipFile(path, "r") as zf:
                for name in sorted(zf.namelist()):
                    if not name.startswith("word/media/"):
                        continue
                    if Path(name).suffix.lower() not in _IMAGE_EXTENSIONS:
                        continue
                    data = zf.read(name)
                    if data:
                        images.append((Path(name).name, data))
        except (zipfile.BadZipFile, OSError):
            logger.warning("Failed to extract images from DOCX: %s", path.name)
        return images

    async def _extract_doc(self, path: Path) -> str:
        """Extract text from legacy DOC via antiword CLI."""
        proc = await asyncio.create_subprocess_exec(
            "antiword",
            "-m",
            "UTF-8.txt",
            str(path),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)

        if proc.returncode != 0:
            err_msg = stderr.decode("utf-8", errors="replace").strip()
            raise ValueError(f"DOC extraction failed: {err_msg}")

        return stdout.decode("utf-8", errors="replace")

    async def _ocr_images(self, images: list[tuple[str, bytes]]) -> tuple[str, dict]:
        """Run OCR on extracted images and return combined text + metadata."""
        if not images:
            return "", {}

        from app.services.ocr_service import OCRService

        ocr = OCRService()
        texts: list[str] = []
        errors = 0

        for filename, data in images:
            mime = _MIME_MAP.get(Path(filename).suffix.lower(), "image/png")
            try:
                result = await ocr.extract_text(data, mime)
                if result.text and result.text.strip():
                    texts.append(result.text.strip())
            except Exception:
                logger.warning("OCR failed for embedded image: %s", filename)
                errors += 1

        meta: dict = {}
        if texts:
            meta["ocr"] = True
            meta["ocr_image_count"] = len(texts)
        if errors:
            meta["ocr_errors"] = errors
        return "\n\n".join(texts), meta
