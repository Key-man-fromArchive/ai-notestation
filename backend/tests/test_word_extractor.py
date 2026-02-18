"""Tests for Word document (.docx/.doc) text extraction."""

from __future__ import annotations

from unittest.mock import AsyncMock, patch

import pytest

from app.services.word_extractor import WordExtractionResult, WordExtractor

# --- WordExtractionResult model tests ---


def test_result_model_basic():
    result = WordExtractionResult(text="hello", page_count=1, metadata={"format": "docx"})
    assert result.text == "hello"
    assert result.page_count == 1
    assert result.metadata == {"format": "docx"}


def test_result_model_doc_format():
    result = WordExtractionResult(text="content", page_count=1, metadata={"format": "doc"})
    assert result.metadata["format"] == "doc"


# --- DOCX extraction tests (python-docx) ---


@pytest.mark.asyncio
async def test_extract_docx_paragraphs(tmp_path):
    """DOCX extraction reads paragraphs."""
    from docx import Document

    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("연구 개요")
    doc.add_paragraph("바이오 디지털 트윈")
    doc.save(str(docx_file))

    extractor = WordExtractor()
    result = await extractor.extract(str(docx_file))

    assert "연구 개요" in result.text
    assert "바이오 디지털 트윈" in result.text
    assert result.metadata["format"] == "docx"
    assert result.page_count == 1


@pytest.mark.asyncio
async def test_extract_docx_with_table(tmp_path):
    """DOCX extraction reads table cells."""
    from docx import Document

    docx_file = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Table test")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value 1"
    table.cell(1, 1).text = "Value 2"
    doc.save(str(docx_file))

    extractor = WordExtractor()
    result = await extractor.extract(str(docx_file))

    assert "Table test" in result.text
    assert "Header A" in result.text
    assert "Value 1" in result.text


@pytest.mark.asyncio
async def test_extract_docx_empty(tmp_path):
    """ValueError for DOCX with no text."""
    from docx import Document

    docx_file = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(docx_file))

    extractor = WordExtractor()
    with pytest.raises(ValueError, match="no extractable text"):
        await extractor.extract(str(docx_file))


# --- DOCX image extraction tests ---

# 1x1 red PNG (smallest valid PNG)
_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00"
    b"\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00"
    b"\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_extract_docx_images(tmp_path):
    """DOCX image extraction picks up images from word/media/ folder."""
    import zipfile

    from docx import Document

    # Create a valid DOCX first, then inject an image into the ZIP
    docx_file = tmp_path / "with_images.docx"
    doc = Document()
    doc.add_paragraph("Has images")
    doc.save(str(docx_file))

    # Append an image to the ZIP
    with zipfile.ZipFile(docx_file, "a") as zf:
        zf.writestr("word/media/image1.png", _TINY_PNG)

    extractor = WordExtractor()
    images = extractor._extract_docx_images(docx_file)

    assert len(images) == 1
    assert images[0][0] == "image1.png"


def test_extract_docx_images_no_media(tmp_path):
    """DOCX without word/media/ returns empty list."""
    from docx import Document

    docx_file = tmp_path / "no_images.docx"
    doc = Document()
    doc.add_paragraph("text only")
    doc.save(str(docx_file))

    extractor = WordExtractor()
    images = extractor._extract_docx_images(docx_file)
    assert images == []


# --- DOC extraction tests (antiword, mocked) ---


@pytest.mark.asyncio
async def test_extract_doc_antiword(tmp_path):
    """DOC extraction delegates to antiword CLI."""
    doc_file = tmp_path / "test.doc"
    doc_file.write_bytes(b"fake-doc-binary")

    extractor = WordExtractor()

    mock_proc = AsyncMock()
    mock_proc.communicate = AsyncMock(return_value=(b"Extracted text from DOC", b""))
    mock_proc.returncode = 0

    with patch("asyncio.create_subprocess_exec", return_value=mock_proc):
        result = await extractor.extract(str(doc_file))

    assert "Extracted text from DOC" in result.text
    assert result.metadata["format"] == "doc"


@pytest.mark.asyncio
async def test_extract_doc_antiword_failure(tmp_path):
    """DOC extraction fails if antiword returns non-zero."""
    doc_file = tmp_path / "bad.doc"
    doc_file.write_bytes(b"corrupt")

    extractor = WordExtractor()

    mock_proc = AsyncMock()
    mock_proc.communicate = AsyncMock(return_value=(b"", b"antiword: file is corrupt"))
    mock_proc.returncode = 1

    with (
        patch("asyncio.create_subprocess_exec", return_value=mock_proc),
        pytest.raises(ValueError, match="DOC extraction failed"),
    ):
        await extractor.extract(str(doc_file))


# --- Error cases ---


@pytest.mark.asyncio
async def test_extract_file_not_found():
    """FileNotFoundError for missing files."""
    extractor = WordExtractor()
    with pytest.raises(FileNotFoundError, match="Word file not found"):
        await extractor.extract("/nonexistent/test.docx")


@pytest.mark.asyncio
async def test_extract_unsupported_extension(tmp_path):
    """ValueError for unsupported extensions."""
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("hello")

    extractor = WordExtractor()
    with pytest.raises(ValueError, match="Unsupported format"):
        await extractor.extract(str(txt_file))


# --- API routing test ---


def test_word_extension_accepted_by_routing():
    """Word extensions are recognized in the files API routing logic."""
    from app.api.files import _WORD_EXTENSIONS

    assert ".docx" in _WORD_EXTENSIONS
    assert ".doc" in _WORD_EXTENSIONS


# --- OCR integration tests ---


def _make_ocr_result(text: str):
    """Create a mock OCRResult."""
    from app.services.ocr_service import OCRResult

    return OCRResult(text=text, confidence=0.9, method="mock")


@pytest.mark.asyncio
async def test_ocr_images_combines_text():
    """OCR results from multiple images are combined."""
    extractor = WordExtractor()
    images = [
        ("img1.png", _TINY_PNG),
        ("img2.jpg", b"\xff\xd8\xff\xe0fake"),
    ]

    mock_extract = AsyncMock(
        side_effect=[
            _make_ocr_result("text from image 1"),
            _make_ocr_result("text from image 2"),
        ]
    )

    with patch("app.services.ocr_service.OCRService.extract_text", mock_extract):
        text, meta = await extractor._ocr_images(images)

    assert "text from image 1" in text
    assert "text from image 2" in text
    assert meta["ocr"] is True
    assert meta["ocr_image_count"] == 2


@pytest.mark.asyncio
async def test_ocr_images_empty_list():
    """No images -> empty text and no OCR metadata."""
    extractor = WordExtractor()
    text, meta = await extractor._ocr_images([])
    assert text == ""
    assert meta == {}


@pytest.mark.asyncio
async def test_extract_docx_text_only_no_ocr_metadata(tmp_path):
    """DOCX with text but no images has no OCR keys in metadata."""
    from docx import Document

    docx_file = tmp_path / "text_only.docx"
    doc = Document()
    doc.add_paragraph("pure text document")
    doc.save(str(docx_file))

    extractor = WordExtractor()
    result = await extractor.extract(str(docx_file))

    assert result.text == "pure text document"
    assert "ocr" not in result.metadata
    assert "ocr_image_count" not in result.metadata


@pytest.mark.asyncio
async def test_extract_docx_with_ocr_integration(tmp_path):
    """Full DOCX extraction: body text + OCR from embedded image."""
    import zipfile

    from docx import Document

    docx_file = tmp_path / "full.docx"
    doc = Document()
    doc.add_paragraph("body text")
    doc.save(str(docx_file))

    # Inject an image into the ZIP
    with zipfile.ZipFile(docx_file, "a") as zf:
        zf.writestr("word/media/scan.png", _TINY_PNG)

    mock_extract = AsyncMock(return_value=_make_ocr_result("scanned text"))

    extractor = WordExtractor()
    with patch("app.services.ocr_service.OCRService.extract_text", mock_extract):
        result = await extractor.extract(str(docx_file))

    assert "body text" in result.text
    assert "scanned text" in result.text
    assert result.metadata["ocr"] is True
    assert result.metadata["format"] == "docx"
